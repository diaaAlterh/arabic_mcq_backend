import os
import json
from typing import List, Dict, Any
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from crewai import Agent, Task, Crew, Process
from crewai.llm import LLM
import pytesseract
from PIL import Image, ImageEnhance, ImageFilter
import cv2
import numpy as np
import re
import PyPDF2
import fitz

# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"  # path
pytesseract.pytesseract.tesseract_cmd = r"/usr/bin/tesseract"  # path

class ArabicMCQGeneratorSystem:
    def __init__(self):
        self.llm = LLM(
            model="gemini/gemini-2.0-flash",
            api_key=os.getenv("GEMINI_API_KEY"),
            temperature=0.7
        )
        self.ocr_config = '--oem 3 --psm 6 -l ara+eng'
        self.setup_agents()

    def setup_agents(self):
        self.mcq_generator_agent = Agent(
            role="خبير توليد أسئلة الاختيار من متعدد",
            goal="إنشاء أسئلة اختيار من متعدد دقيقة ومتنوعة من النصوص العربية",
            backstory="أنت خبير تربوي متخصص في إنشاء أسئلة الاختيار من متعدد باللغة العربية.",
            verbose=True, 
            allow_delegation=False, 
            llm=self.llm
        )
        
        self.mcq_validator_agent = Agent(
            role="مدقق صحة أسئلة الاختيار من متعدد",
            goal="التأكد من صحة ودقة أسئلة الاختيار من متعدد والخيارات المولدة",
            backstory="أنت مدقق محترف لأسئلة الاختيار من متعدد باللغة العربية.",
            verbose=True, 
            allow_delegation=False, 
            llm=self.llm
        )
        
        self.mcq_difficulty_agent = Agent(
            role="خبير تصنيف مستوى صعوبة أسئلة الاختيار من متعدد",
            goal="تصنيف أسئلة الاختيار من متعدد حسب مستوى الصعوبة (سهل، متوسط، صعب)",
            backstory="أنت خبير في علم النفس التربوي وتقييم أسئلة الاختيار من متعدد.",
            verbose=True, 
            allow_delegation=False, 
            llm=self.llm
        )
        
        self.options_quality_agent = Agent(
            role="فاحص جودة خيارات الاختيار من متعدد",
            goal="ضمان جودة وتنوع خيارات أسئلة الاختيار من متعدد",
            backstory="أنت خبير في تقييم جودة خيارات أسئلة الاختيار من متعدد.",
            verbose=True, 
            allow_delegation=False, 
            llm=self.llm
        )

    def preprocess_image_for_ocr(self, path: str) -> np.ndarray:
        img = cv2.imread(path)
        if img is None:
            raise ValueError("لا يمكن قراءة الصورة")
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        denoised = cv2.medianBlur(gray, 3)
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        enhanced = clahe.apply(denoised)
        kernel = np.array([[-1, -1, -1], [-1, 9, -1], [-1, -1, -1]])
        sharp = cv2.filter2D(enhanced, -1, kernel)
        _, bw = cv2.threshold(sharp, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        clean = cv2.morphologyEx(bw, cv2.MORPH_CLOSE, np.ones((2, 2), np.uint8))
        return clean

    def extract_text_from_image(self, path: str) -> str:
        try:
            im = self.preprocess_image_for_ocr(path)
            pil_img = Image.fromarray(im)
            pil_img = ImageEnhance.Contrast(pil_img).enhance(1.5)
            pil_img = pil_img.filter(ImageFilter.MedianFilter(size=3))
            txt = pytesseract.image_to_string(pil_img, config='--oem 1 --psm 3', lang='ara+eng')
            cleaned = self.clean_extracted_text(txt)
            if not cleaned.strip():
                txt = pytesseract.image_to_string(
                    pil_img, config='--oem 1 --psm 3', lang='ara+eng'
                )
                cleaned = self.clean_extracted_text(txt)
            return cleaned or "لم يتم العثور على نص في الصورة"
        except Exception as e:
            return f"خطأ في استخراج النص: {str(e)}"

    def clean_extracted_text(self, text: str) -> str:
        t = re.sub(
            r'[^\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF'
            r'\s\w\d\.,!?:;\-()]',
            '', text
        )
        t = re.sub(r'\n\s*\n', '\n', t)
        return re.sub(r'\s+', ' ', t).strip()

    def get_text_confidence(self, path: str) -> Dict[str, Any]:
        try:
            im = self.preprocess_image_for_ocr(path)
            pil_img = Image.fromarray(im)
            d = pytesseract.image_to_data(
                pil_img, config=self.ocr_config, lang='ara+eng', output_type=pytesseract.Output.DICT
            )
            confs = [int(c) for c in d['conf'] if c.isdigit() and int(c) > 0]
            return {
                'average_confidence': sum(confs) / len(confs) if confs else 0,
                'word_count': len([w for w in d['text'] if w.strip()]),
                'high_confidence_words': len([c for c in confs if c > 70])
            }
        except Exception:
            return {'average_confidence': 0, 'word_count': 0, 'high_confidence_words': 0}

    def extract_text_from_pdf(self, path: str, page_range: str = "all") -> str:
        try:
            text = ""
            
            try:
                pdf_document = fitz.open(path)
                total_pages = len(pdf_document)
                
                if page_range.lower() == "all":
                    pages_to_extract = range(total_pages)
                else:
                    pages_to_extract = self.parse_page_range(page_range, total_pages)
                
                for page_num in pages_to_extract:
                    if 0 <= page_num < total_pages:
                        page = pdf_document[page_num]
                        page_text = page.get_text()
                        text += page_text + "\n"
                
                pdf_document.close()
                
            except Exception as e:
                print(f"فشل PyMuPDF، جاري المحاولة مع PyPDF2: {e}")
                
                with open(path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    total_pages = len(pdf_reader.pages)
                    
                    if page_range.lower() == "all":
                        pages_to_extract = range(total_pages)
                    else:
                        pages_to_extract = self.parse_page_range(page_range, total_pages)
                    
                    for page_num in pages_to_extract:
                        if 0 <= page_num < total_pages:
                            page = pdf_reader.pages[page_num]
                            text += page.extract_text() + "\n"
            
            cleaned_text = self.clean_extracted_text(text)
            return cleaned_text if cleaned_text.strip() else "لم يتم العثور على نص في PDF"
            
        except Exception as e:
            return f"خطأ في قراءة PDF: {str(e)}"

    def parse_page_range(self, page_range: str, total_pages: int) -> List[int]:
        pages = []
        
        try:
            parts = page_range.split(',')
            
            for part in parts:
                part = part.strip()
                if '-' in part:
                    start, end = part.split('-')
                    start = int(start.strip()) - 1
                    end = int(end.strip()) - 1
                    pages.extend(range(start, min(end + 1, total_pages)))
                else:
                    page_num = int(part.strip()) - 1
                    if 0 <= page_num < total_pages:
                        pages.append(page_num)
            
            return sorted(list(set(pages)))
            
        except Exception as e:
            print(f"خطأ في تحليل نطاق الصفحات: {e}")
            return list(range(min(5, total_pages)))

    def get_pdf_info(self, path: str) -> Dict[str, Any]:
        try:
            pdf_document = fitz.open(path)
            total_pages = len(pdf_document)
            
            sample_text = ""
            for i in range(min(3, total_pages)):
                page = pdf_document[i]
                sample_text += page.get_text()[:500]
            
            pdf_document.close()
            
            has_arabic = bool(re.search(r'[\u0600-\u06FF]', sample_text))
            
            return {
                'total_pages': total_pages,
                'has_arabic': has_arabic,
                'sample_text_length': len(sample_text.strip()),
                'estimated_total_chars': len(sample_text) * total_pages // min(3, total_pages) if total_pages > 0 else 0
            }
            
        except Exception as e:
            return {
                'total_pages': 0,
                'has_arabic': False,
                'error': str(e),
                'sample_text_length': 0,
                'estimated_total_chars': 0
            }

    def create_mcq_tasks(self, text: str, n: int = 5) -> List[Task]:
        generation_description = f"""
قم بتوليد {n} أسئلة اختيار من متعدد من النص التالي:

{text}

يجب أن تكون النتيجة في صيغة JSON بالشكل التالي:
{{
    "questions": [
        {{
            "question": "نص السؤال",
            "options": ["الخيار أ", "الخيار ب", "الخيار ج", "الخيار د"],
            "correct_answer": "أ",
            "explanation": "تفسير الإجابة الصحيحة",
            "difficulty": "متوسط"
        }}
    ]
}}

تأكد من:
- كل سؤال له 4 خيارات بالضبط
- إجابة واحدة صحيحة فقط
- الخيارات الخاطئة منطقية ومعقولة
"""

        validation_description = """
راجع الأسئلة المولدة وتأكد من:
1. وضوح السؤال وعدم الغموض
2. وجود إجابة واحدة صحيحة فقط
3. منطقية الخيارات الخاطئة
4. صحة التفسير
أرجع النتيجة بنفس تنسيق JSON المطلوب.
"""

        difficulty_description = """
صنف كل سؤال من الأسئلة المولدة حسب مستوى الصعوبة:
- سهل: أسئلة تذكر واستدعاء معلومات مباشرة
- متوسط: أسئلة فهم وتطبيق
- صعب: أسئلة تحليل وتقييم ونقد

أضف أو حدث حقل "difficulty" لكل سؤال في JSON.
"""

        options_description = """
راجع جودة خيارات كل سؤال وتأكد من:
1. الخيارات الخاطئة منطقية وليست واضحة الخطأ
2. تنوع الخيارات ومناسبتها للسؤال
3. عدم وجود خيارات متشابهة أو مكررة
4. توازن طول الخيارات

حسن الخيارات إذا لزم الأمر وأرجع JSON النهائي.
"""

        return [
            Task(
                description=generation_description,
                agent=self.mcq_generator_agent,
                expected_output="JSON format with questions array"
            ),
            Task(
                description=validation_description,
                agent=self.mcq_validator_agent,
                expected_output="Validated JSON format with questions array"
            ),
            Task(
                description=difficulty_description,
                agent=self.mcq_difficulty_agent,
                expected_output="JSON with difficulty classification"
            ),
            Task(
                description=options_description,
                agent=self.options_quality_agent,
                expected_output="Final improved JSON format"
            )
        ]

    def parse_crew_output(self, crew_output) -> Dict[str, Any]:
        try:
            if hasattr(crew_output, 'raw'):
                output_text = crew_output.raw
            elif hasattr(crew_output, 'content'):
                output_text = crew_output.content
            elif hasattr(crew_output, 'result'):
                output_text = crew_output.result
            elif isinstance(crew_output, str):
                output_text = crew_output
            else:
                output_text = str(crew_output)

            print(f"محاولة تحليل المخرجات: {output_text[:200]}...")

            json_match = re.search(r'\{.*\}', output_text, re.DOTALL)
            if json_match:
                json_str = json_match.group()
                try:
                    return json.loads(json_str)
                except json.JSONDecodeError as e:
                    print(f"خطأ في تحليل JSON: {e}")
                    cleaned_json = self.clean_json_string(json_str)
                    return json.loads(cleaned_json)

            return self.create_fallback_questions(output_text)

        except Exception as e:
            print(f"خطأ في تحليل المخرجات: {e}")
            return self.create_sample_questions()

    def clean_json_string(self, json_str: str) -> str:
        json_str = re.sub(r'//.*?\n', '', json_str)
        json_str = re.sub(r'/\*.*?\*/', '', json_str, flags=re.DOTALL)
        json_str = re.sub(r'"\s*\n\s*"', '",\n"', json_str)
        json_str = re.sub(r'}\s*\n\s*{', '},\n{', json_str)
        return json_str

    def create_fallback_questions(self, output_text: str) -> Dict[str, Any]:
        questions = []
        lines = output_text.split('\n')
        current_question = {}
        
        for line in lines:
            line = line.strip()
            if 'سؤال' in line or line.startswith('Q') or line.startswith('السؤال'):
                if current_question and 'question' in current_question:
                    questions.append(current_question)
                current_question = {'question': line, 'options': [], 'correct_answer': 'أ', 'explanation': '', 'difficulty': 'متوسط'}
            elif line.startswith(('أ', 'ب', 'ج', 'د', 'A', 'B', 'C', 'D')):
                if 'options' in current_question:
                    current_question['options'].append(line)
        
        if current_question and 'question' in current_question:
            questions.append(current_question)
            
        if not questions:
            return self.create_sample_questions()
            
        return {'questions': questions}

    def create_sample_questions(self) -> Dict[str, Any]:
        return {
            'questions': [
                {
                    'question': 'سؤال تجريبي - ما هي عاصمة المملكة العربية السعودية؟',
                    'options': ['الرياض', 'جدة', 'الدمام', 'مكة المكرمة'],
                    'correct_answer': 'أ',
                    'explanation': 'الرياض هي العاصمة الرسمية للمملكة العربية السعودية',
                    'difficulty': 'سهل'
                }
            ]
        }

    def save_mcq_to_word(self, data: Dict[str, Any], filename: str):
        try:
            doc = Document()
            
            header = doc.add_heading('أسئلة الاختيار من متعدد', 0)
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            info_para = doc.add_paragraph(
                f"تاريخ الإنشاء: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
                f"عدد الأسئلة: {len(data.get('questions', []))}\n"
            )
            info_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            doc.add_page_break()

            questions = data.get('questions', [])
            for i, question in enumerate(questions, 1):
                question_para = doc.add_paragraph()
                question_run = question_para.add_run(f"السؤال {i}: {question.get('question', '')}")
                question_run.bold = True
                question_run.font.size = Pt(14)
                question_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                options = question.get('options', [])
                labels = ['أ', 'ب', 'ج', 'د']
                
                for idx, option in enumerate(options[:4]):
                    option_para = doc.add_paragraph(f"    {labels[idx]}. {option}")
                    option_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    option_para.runs[0].font.size = Pt(12)

                answer_para = doc.add_paragraph()
                answer_run = answer_para.add_run(f"الإجابة الصحيحة: {question.get('correct_answer', 'أ')}")
                answer_run.bold = True
                answer_run.font.color.rgb = RGBColor(0, 128, 0)
                answer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                difficulty = question.get('difficulty', 'متوسط')
                if difficulty:
                    diff_para = doc.add_paragraph(f"مستوى الصعوبة: {difficulty}")
                    diff_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    if difficulty == 'سهل':
                        color = RGBColor(0, 128, 0)
                    elif difficulty == 'متوسط':
                        color = RGBColor(255, 140, 0)
                    else:
                        color = RGBColor(255, 0, 0)
                    
                    diff_para.runs[0].font.color.rgb = color

                explanation = question.get('explanation', '')
                if explanation:
                    expl_para = doc.add_paragraph(f"التفسير: {explanation}")
                    expl_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    expl_para.runs[0].font.size = Pt(11)

                separator = doc.add_paragraph("─" * 50)
                separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()

            doc.save(filename)
            with open(filename.replace('.docx', '.json'), 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            print(f"✔️ تم حفظ الملف بنجاح: {filename}")
            print(f"✔️ تم حفظ الملف بنجاح: {filename.replace('.docx', '.json')}")

            return True
            
        except Exception as e:
            print(f"❌ خطأ في حفظ الملف: {e}")
            return False

    def show_questions_preview(self, questions: List[Dict]):
        print("\n" + "="*50)
        print("معاينة الأسئلة")
        print("="*50)
        
        for i, q in enumerate(questions, 1):
            print(f"\nالسؤال {i}: {q.get('question', '')}")
            
            options = q.get('options', [])
            labels = ['أ', 'ب', 'ج', 'د']
            for idx, option in enumerate(options[:4]):
                print(f"  {labels[idx]}. {option}")
            
            print(f"الإجابة الصحيحة: {q.get('correct_answer', '')}")
            
            if q.get('difficulty'):
                print(f"الصعوبة: {q.get('difficulty')}")
                
            if q.get('explanation'):
                print(f"التفسير: {q.get('explanation')}")
            
            print("-" * 30)

    def run_mcq_generation(self, text: str, mode: str, n: int, **kwargs) -> Dict[str, Any]:
        try:
            if mode == 'image':
                print("🔍 جاري استخراج النص من الصورة...")
                confidence = self.get_text_confidence(text)
                print(f"مستوى ثقة OCR: {confidence['average_confidence']:.1f}%")
                
                extracted_text = self.extract_text_from_image(text)
                if extracted_text.startswith("خطأ"):
                    return {'error': extracted_text, 'confidence': confidence}
                    
                print(f"تم استخراج النص بنجاح ({len(extracted_text)} حرف)")
                final_text = extracted_text
                
            elif mode == 'pdf':
                print("📄 جاري استخراج النص من PDF...")
                page_range = kwargs.get('page_range', 'all')
                extracted_text = self.extract_text_from_pdf(text, page_range)
                
                if extracted_text.startswith("خطأ"):
                    return {'error': extracted_text}
                    
                print(f"تم استخراج النص بنجاح ({len(extracted_text)} حرف)")
                final_text = extracted_text
                
            else:
                final_text = text

            if len(final_text.strip()) < 50:
                return {'error': 'النص قصير جداً لتوليد أسئلة مفيدة (يجب أن يكون أكثر من 50 حرف)'}

            print(f"🎯 جاري توليد {n} أسئلة...")
            
            tasks = self.create_mcq_tasks(final_text, n)
            
            crew = Crew(
                agents=[
                    self.mcq_generator_agent, 
                    self.mcq_validator_agent,
                    self.mcq_difficulty_agent,
                    self.options_quality_agent
                ],
                tasks=tasks,
                process=Process.sequential,
                verbose=True
            )
            
            result = crew.kickoff()
            parsed_result = self.parse_crew_output(result)
            
            print("✔️ تم توليد الأسئلة بنجاح!")
            return parsed_result
            
        except Exception as e:
            print(f"❌ خطأ في توليد الأسئلة: {e}")
            return {'error': f'خطأ في التوليد: {str(e)}'}

    def create_interactive_menu(self):
        print("🎯 مرحباً بك في نظام توليد أسئلة الاختيار من متعدد")
        print("="*60)
        
        while True:
            print("\n" + "="*40)
            print("الخيارات المتاحة:")
            print("1) إدخال نص مباشر")
            print("2) قراءة من ملف نصي")
            print("3) استخراج من صورة")
            print("4) استخراج من ملف PDF")
            print("0) خروج")
            print("="*40)
            
            choice = input("اختر رقماً: ").strip()
            
            if choice == '1':
                self.handle_text_input()
            elif choice == '2':
                self.handle_file_input()
            elif choice == '3':
                self.handle_image_input()
            elif choice == '4':
                self.handle_pdf_input()
            elif choice == '0':
                print("شكراً لاستخدام النظام، مع السلامة! 👋")
                break
            else:
                print("❌ خيار غير صالح، حاول مرة أخرى")

    def handle_text_input(self):
        print("\n📝 إدخال النص:")
        print("(أدخل النص واضغط Enter مرتين للانتهاء)")
        
        lines = []
        empty_lines = 0
        
        while empty_lines < 2:
            line = input()
            if line.strip() == "":
                empty_lines += 1
            else:
                empty_lines = 0
                lines.append(line)
        
        text = "\n".join(lines).strip()
        
        if len(text) < 50:
            print("❌ النص قصير جداً، يجب أن يكون أكثر من 50 حرف")
            return
            
        n = self.get_question_count()
        self.process_content(text, 'text', n)

    def handle_file_input(self):
        file_path = input("\n📁 أدخل مسار الملف النصي: ").strip()
        
        if not os.path.exists(file_path):
            print("❌ الملف غير موجود")
            return
            
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read().strip()
                
            if len(text) < 50:
                print("❌ محتوى الملف قصير جداً")
                return
                
            print(f"✔️ تم قراءة الملف بنجاح ({len(text)} حرف)")
            n = self.get_question_count()
            self.process_content(text, 'text', n)
            
        except Exception as e:
            print(f"❌ خطأ في قراءة الملف: {e}")

    def handle_image_input(self):
        image_path = input("\n🖼️ أدخل مسار الصورة: ").strip()
        
        # إزالة علامات الاقتباس إن وجدت
        image_path = image_path.strip('"').strip("'")
        
        if not image_path:
            print("❌ يرجى إدخال مسار صحيح للصورة")
            return
            
        if not os.path.exists(image_path):
            print("❌ الصورة غير موجودة في المسار المحدد")
            return
        
        # التحقق من صيغة الصورة
        valid_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp']
        if not any(image_path.lower().endswith(ext) for ext in valid_extensions):
            print("❌ صيغة الصورة غير مدعومة. الصيغ المدعومة: JPG, PNG, GIF, BMP, TIFF, WEBP")
            return
        
        try:
            n = self.get_question_count()
            print(f"🔄 جاري معالجة الصورة وتوليد {n} أسئلة...")
            self.process_content(image_path, 'image', n)
        except Exception as e:
            print(f"❌ خطأ أثناء معالجة الصورة: {str(e)}")

    def handle_pdf_input(self):
        """معالجة إدخال ملفات PDF مع التحقق الشامل"""
        pdf_path = input("\n📄 أدخل مسار ملف PDF: ").strip()
        
        # إزالة علامات الاقتباس إن وجدت
        pdf_path = pdf_path.strip('"').strip("'")
        
        if not pdf_path:
            print("❌ يرجى إدخال مسار صحيح لملف PDF")
            return
            
        if not os.path.exists(pdf_path):
            print("❌ الملف غير موجود في المسار المحدد")
            return
            
        if not pdf_path.lower().endswith('.pdf'):
            print("❌ يجب أن يكون الملف بصيغة PDF")
            return
        
        # التحقق من حجم الملف
        try:
            file_size = os.path.getsize(pdf_path) / (1024 * 1024)  # بالميجابايت
            if file_size > 50:  # حد أقصى 50 ميجا
                print(f"⚠️ حجم الملف كبير ({file_size:.1f} MB). قد يستغرق وقتاً أطول للمعالجة")
        except:
            pass
        
        print("\n🔍 خيارات استخراج الصفحات:")
        print("1️⃣ جميع الصفحات (all)")
        print("2️⃣ صفحات محددة (مثال: 1-5 أو 1,3,5-8)")
        
        page_range = input("اختر نطاق الصفحات: ").strip()
        
        if not page_range:
            page_range = "all"
        
        try:
            n = self.get_question_count()
            print(f"🔄 جاري معالجة PDF وتوليد {n} أسئلة...")
            self.process_content(pdf_path, 'pdf', n, page_range=page_range)
        except Exception as e:
            print(f"❌ خطأ أثناء معالجة PDF: {str(e)}")

    def get_question_count(self) -> int:
        """الحصول على عدد الأسئلة المطلوب توليدها مع التحقق من صحة الإدخال"""
        while True:
            try:
                user_input = input("🔢 كم عدد الأسئلة التي تريد توليدها؟ (افتراضي 5، الحد الأقصى 50): ").strip()
                
                if not user_input:
                    return 5  # القيمة الافتراضية
                
                n = int(user_input)
                
                if n < 1:
                    print("❌ يجب أن يكون العدد أكبر من صفر")
                    continue
                elif n > 50:
                    print("❌ الحد الأقصى هو 50 سؤال")
                    continue
                
                return n
                
            except ValueError:
                print("❌ يرجى إدخال رقم صحيح")
                continue

    def process_content(self, source: str, mode: str, n: int, **kwargs):
        """معالجة المحتوى وتوليد الأسئلة مع معالجة شاملة للأخطاء"""
        try:
            # إظهار شريط التقدم أو رسالة الانتظار
            print("⏳ جاري المعالجة، يرجى الانتظار...")
            
            # تشغيل توليد الأسئلة
            result = self.run_mcq_generation(source, mode, n, **kwargs)
            
            if "error" in result:
                print(f"❌ خطأ في المعالجة: {result['error']}")
                return
            
            if not result or "questions" not in result:
                print("❌ فشل في توليد الأسئلة. يرجى المحاولة مرة أخرى")
                return
            
            # إنشاء اسم الملف مع الطابع الزمني
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            if mode in ['image', 'pdf']:
                source_name = os.path.splitext(os.path.basename(source))[0]
            else:
                source_name = "text_input"
            filename = f"mcq_{source_name}_{timestamp}.docx"
            
            print(f"✅ تم توليد {len(result.get('questions', []))} سؤال بنجاح")
            print(f"💾 جاري حفظ النتائج في ملف: {filename}")
            
            # حفظ النتائج
            success = self.save_mcq_to_word(result, filename)
            
            if success:
                print(f"✅ تم حفظ الملف بنجاح: {filename}")
                self.post_generation_options(result, filename)
            else:
                print("❌ فشل في حفظ الملف")
                
        except FileNotFoundError:
            print("❌ الملف المحدد غير موجود")
        except PermissionError:
            print("❌ ليس لديك صلاحية للوصول إلى الملف")
        except Exception as e:
            print(f"❌ خطأ غير متوقع: {str(e)}")
            print("🔄 يرجى المحاولة مرة أخرى أو التحقق من المسار")

    def post_generation_options(self, result: dict, filename: str):
        """خيارات ما بعد توليد الأسئلة"""
        print("\n" + "="*50)
        print("🎉 تم إنجاز المهمة بنجاح!")
        print("="*50)
        
        while True:
            print("\n📋 الخيارات المتاحة:")
            print("1️⃣ عرض الأسئلة في المحطة الطرفية")
            print("2️⃣ فتح ملف Word")
            print("3️⃣ إنشاء ملف HTML للعرض في المتصفح")
            print("4️⃣ العودة للقائمة الرئيسية")
            print("5️⃣ إنهاء البرنامج")
            
            choice = input("\nاختر خياراً (1-5): ").strip()
            
            if choice == '1':
                self.display_questions_in_terminal(result)
            elif choice == '2':
                self.open_word_file(filename)
            elif choice == '3':
                self.create_html_output(result, filename)
            elif choice == '4':
                break
            elif choice == '5':
                print("👋 شكراً لاستخدام المولد!")
                exit()
            else:
                print("❌ خيار غير صحيح، يرجى اختيار رقم من 1 إلى 5")

    def display_questions_in_terminal(self, result: dict):
        """عرض الأسئلة في المحطة الطرفية بتنسيق جميل"""
        questions = result.get('questions', [])
        if not questions:
            print("❌ لا توجد أسئلة للعرض")
            return
        
        print("\n" + "="*60)
        print("📚 الأسئلة المُولدة")
        print("="*60)
        
        for i, q in enumerate(questions, 1):
            print(f"\n🔹 السؤال {i}:")
            print(f"❓ {q.get('question', 'غير متوفر')}")
            
            options = q.get('options', [])
            labels = ['أ', 'ب', 'ج', 'د']
            correct_answer = q.get('correct_answer', 'أ')
            
            for j, option in enumerate(options[:4]):
                marker = "✅" if labels[j] == correct_answer else "⚪"
                print(f"   {marker} {labels[j]}. {option}")
            
            if 'explanation' in q and q['explanation']:
                print(f"💡 التفسير: {q['explanation']}")
            
            if 'difficulty' in q and q['difficulty']:
                print(f"📊 مستوى الصعوبة: {q['difficulty']}")
            
            print("-" * 40)
        
        input("\nاضغط Enter للمتابعة...")

    def open_word_file(self, filename: str):
        """فتح ملف Word باستخدام البرنامج الافتراضي"""
        try:
            if os.path.exists(filename):
                import subprocess
                import platform
                
                if platform.system() == 'Windows':
                    os.startfile(filename)
                elif platform.system() == 'Darwin':  # macOS
                    subprocess.run(['open', filename])
                else:  # Linux
                    subprocess.run(['xdg-open', filename])
                
                print(f"✅ تم فتح الملف: {filename}")
            else:
                print("❌ الملف غير موجود")
        except Exception as e:
            print(f"❌ فشل في فتح الملف: {str(e)}")

    def create_html_output(self, result: dict, base_filename: str):
        """إنشاء ملف HTML لعرض الأسئلة في المتصفح"""
        try:
            html_filename = base_filename.replace('.docx', '.html')
            questions = result.get('questions', [])
            
            html_content = f"""
<!DOCTYPE html>
<html dir="rtl" lang="ar">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>أسئلة الاختيار من متعدد</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            margin: 0;
            padding: 20px;
            background-color: #f5f5f5;
        }}
        .container {{
            max-width: 800px;
            margin: 0 auto;
            background: white;
            padding: 30px;
            border-radius: 10px;
            box-shadow: 0 0 20px rgba(0,0,0,0.1);
        }}
        h1 {{
            color: #2c3e50;
            text-align: center;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
        }}
        .question {{
            margin-bottom: 30px;
            padding: 20px;
            border: 1px solid #ddd;
            border-radius: 8px;
            background: #fafafa;
        }}
        .question-title {{
            font-weight: bold;
            font-size: 18px;
            color: #2c3e50;
            margin-bottom: 15px;
        }}
        .option {{
            margin: 8px 0;
            padding: 10px;
            border-radius: 5px;
            cursor: pointer;
        }}
        .option.correct {{
            background-color: #d4edda;
            border: 2px solid #28a745;
            font-weight: bold;
        }}
        .option.incorrect {{
            background-color: #f8f9fa;
            border: 1px solid #dee2e6;
        }}
        .explanation {{
            margin-top: 15px;
            padding: 15px;
            background-color: #e7f3ff;
            border-right: 4px solid #007bff;
            border-radius: 5px;
        }}
        .difficulty {{
            margin-top: 10px;
            padding: 5px 10px;
            border-radius: 15px;
            display: inline-block;
            font-size: 12px;
            font-weight: bold;
        }}
        .difficulty.easy {{
            background-color: #d4edda;
            color: #155724;
        }}
        .difficulty.medium {{
            background-color: #fff3cd;
            color: #856404;
        }}
        .difficulty.hard {{
            background-color: #f8d7da;
            color: #721c24;
        }}
        .question-number {{
            display: inline-block;
            background: #3498db;
            color: white;
            padding: 5px 10px;
            border-radius: 50%;
            margin-left: 10px;
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>🎓 أسئلة الاختيار من متعدد</h1>
        <p style="text-align: center; color: #666;">عدد الأسئلة: {len(questions)}</p>
"""
            
            for i, q in enumerate(questions, 1):
                html_content += f"""
        <div class="question">
            <div class="question-title">
                <span class="question-number">{i}</span>
                {q.get('question', 'غير متوفر')}
            </div>
"""
                
                options = q.get('options', [])
                labels = ['أ', 'ب', 'ج', 'د']
                correct_answer = q.get('correct_answer', 'أ')
                
                for j, option in enumerate(options[:4]):
                    option_class = "correct" if labels[j] == correct_answer else "incorrect"
                    icon = "✅" if labels[j] == correct_answer else "⚪"
                    html_content += f"""
            <div class="option {option_class}">
                {icon} {labels[j]}. {option}
            </div>
"""
                
                if 'explanation' in q and q['explanation']:
                    html_content += f"""
            <div class="explanation">
                <strong>💡 التفسير:</strong> {q['explanation']}
            </div>
"""
                
                if 'difficulty' in q and q['difficulty']:
                    difficulty = q['difficulty']
                    difficulty_class = ""
                    if difficulty == 'سهل':
                        difficulty_class = "easy"
                    elif difficulty == 'متوسط':
                        difficulty_class = "medium"
                    else:
                        difficulty_class = "hard"
                    
                    html_content += f"""
            <div class="difficulty {difficulty_class}">
                📊 مستوى الصعوبة: {difficulty}
            </div>
"""
                
                html_content += "        </div>\n"
            
            html_content += """
    </div>
</body>
</html>
"""
            
            with open(html_filename, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            print(f"✅ تم إنشاء ملف HTML: {html_filename}")
            
            # فتح الملف في المتصفح
            try:
                import webbrowser
                webbrowser.open(f'file://{os.path.abspath(html_filename)}')
                print("🌐 تم فتح الملف في المتصفح")
            except:
                print("⚠️ لم يتم فتح المتصفح تلقائياً")
                
        except Exception as e:
            print(f"❌ فشل في إنشاء ملف HTML: {str(e)}")


# الجزء الرئيسي من البرنامج
if __name__ == "__main__":
    try:
        # التحقق من المتطلبات الأساسية
        print("🔍 جاري التحقق من المتطلبات...")
        
        # التحقق من وجود API Key
        if not os.getenv("GEMINI_API_KEY"):
            print("❌ يرجى تعيين متغير البيئة GEMINI_API_KEY")
            print("💡 يمكنك الحصول على API Key من: https://makersuite.google.com/app/apikey")
            exit(1)
        
        print("✅ تم العثور على API Key")
        
        # بدء تشغيل النظام
        generator = ArabicMCQGeneratorSystem()
        generator.create_interactive_menu()
        
    except KeyboardInterrupt:
        print("\n\n👋 تم إنهاء البرنامج بواسطة المستخدم")
    except ImportError as e:
        print(f"\n❌ مكتبة مفقودة: {str(e)}")
        print("💡 يرجى تثبيت المكتبات المطلوبة باستخدام:")
        print("pip install crewai python-docx pytesseract pillow opencv-python numpy PyPDF2 pymupdf")
    except Exception as e:
        print(f"\n❌ خطأ في تشغيل البرنامج: {str(e)}")
        print("🔄 يرجى إعادة تشغيل البرنامج")